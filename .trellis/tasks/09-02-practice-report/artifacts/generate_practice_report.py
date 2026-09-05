from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from practice_report_content import (
    ABSTRACT,
    BODY_SECTIONS,
    KEYWORDS,
    REFERENCES,
    REPORT_TITLE,
)


DEFAULT_OUTPUT_PATH = Path(r"C:\Users\Lenovo\Desktop\大学生生成式人工智能应用社会实践报告.docx")


def _set_run_font(run, chinese_font: str, size: Pt, *, bold: bool = False) -> None:
    run.font.name = chinese_font
    run.font.size = size
    run.font.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), chinese_font)


def _set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    _set_run_font(run, "宋体", Pt(12), bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix = paragraph.add_run("第 ")
    _set_run_font(prefix, "宋体", Pt(9))

    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), "宋体")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    properties.extend([fonts, size])
    run.append(properties)
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    field.append(run)
    paragraph._p.append(field)

    suffix = paragraph.add_run(" 页")
    _set_run_font(suffix, "宋体", Pt(9))


def _configure_document(document: DocumentType) -> None:
    for section in document.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)
        _add_page_number(section.footer.paragraphs[0])

    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(0)

    heading_one = document.styles["Heading 1"]
    heading_one.font.name = "黑体"
    heading_one.font.size = Pt(16)
    heading_one.font.bold = True
    heading_one._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "黑体")
    heading_one.paragraph_format.space_before = Pt(12)
    heading_one.paragraph_format.space_after = Pt(6)

    heading_two = document.styles["Heading 2"]
    heading_two.font.name = "黑体"
    heading_two.font.size = Pt(14)
    heading_two.font.bold = True
    heading_two._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "黑体")


def _add_centered_paragraph(
    document: DocumentType,
    text: str,
    *,
    font: str = "黑体",
    size: Pt = Pt(16),
    bold: bool = False,
    space_before: Pt = Pt(0),
    space_after: Pt = Pt(0),
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = space_before
    paragraph.paragraph_format.space_after = space_after
    run = paragraph.add_run(text)
    _set_run_font(run, font, size, bold=bold)


def _add_body_paragraph(document: DocumentType, text: str, *, indent: bool = True) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(0)
    if indent:
        paragraph.paragraph_format.first_line_indent = Pt(24)
    run = paragraph.add_run(text)
    _set_run_font(run, "宋体", Pt(12))


def _add_cover(document: DocumentType) -> None:
    _add_centered_paragraph(
        document,
        "社会调查与实践课程报告",
        size=Pt(22),
        bold=True,
        space_before=Pt(72),
        space_after=Pt(36),
    )
    _add_centered_paragraph(
        document,
        REPORT_TITLE,
        size=Pt(18),
        bold=True,
        space_after=Pt(48),
    )

    fields = [
        ("学校", "____________________________"),
        ("姓名", "____________________________"),
        ("学号", "____________________________"),
        ("专业", "计算机科学与技术"),
        ("班级", "____________________________"),
        ("实践单位", "____________________________"),
        ("实践日期", "____________________________"),
    ]
    table = document.add_table(rows=len(fields), cols=2)
    table.autofit = False
    table.columns[0].width = Cm(4)
    table.columns[1].width = Cm(10)
    for row, (label, value) in zip(table.rows, fields, strict=True):
        row.height = Cm(1)
        _set_cell_text(row.cells[0], label, bold=True)
        _set_cell_text(row.cells[1], value)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(24)
    run = note.add_run("说明：请在提交前据实补全身份信息并核对课程要求。")
    _set_run_font(run, "宋体", Pt(10))


def _add_abstract(document: DocumentType) -> None:
    document.add_page_break()
    heading = document.add_heading("摘要", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_body_paragraph(document, ABSTRACT)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    label = paragraph.add_run("关键词：")
    _set_run_font(label, "黑体", Pt(12), bold=True)
    values = paragraph.add_run("；".join(KEYWORDS))
    _set_run_font(values, "宋体", Pt(12))


def _add_main_body(document: DocumentType) -> None:
    document.add_page_break()
    marker = document.add_paragraph("【正文开始】")
    marker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    marker.runs[0].font.hidden = True

    for heading_text, paragraphs in BODY_SECTIONS:
        document.add_heading(heading_text, level=1)
        for paragraph_text in paragraphs:
            _add_body_paragraph(document, paragraph_text)

    marker = document.add_paragraph("【正文结束】")
    marker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    marker.runs[0].font.hidden = True


def _add_references(document: DocumentType) -> None:
    document.add_page_break()
    document.add_heading("参考资料", level=1)
    for reference in REFERENCES:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph.paragraph_format.first_line_indent = Pt(-24)
        paragraph.paragraph_format.left_indent = Pt(24)
        run = paragraph.add_run(reference)
        _set_run_font(run, "宋体", Pt(10.5))
    _add_body_paragraph(
        document,
        "引用说明：UNESCO 原始条目在当前网络环境未能直接读取，本文仅采用可确认的出版物元数据和一般原则，"
        "未摘录其中的统计数字或逐字引文。",
    )


def _add_process_appendix(document: DocumentType) -> None:
    document.add_page_break()
    document.add_heading("过程材料附页", level=1)
    _add_body_paragraph(
        document,
        "材料状态说明：下列资料检索记录属于现有材料；访谈记录和实践照片属于待补材料。"
        "如课程要求补充实践过程，只能在真实实施后据实填写，不得把提纲或空白粘贴位视为已完成证据。",
    )

    document.add_heading("附页一：资料检索记录（现有材料）", level=2)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["机构", "资料题名", "年份", "用途与核验状态"]
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        _set_cell_text(cell, header, bold=True)
    records = [
        ("教育部", "高等学校人工智能创新行动计划", "2018", "政策背景；官方页面已核验"),
        (
            "国家互联网信息办公室等七部门",
            "生成式人工智能服务管理暂行办法",
            "2023",
            "规范要求；官方页面已核验",
        ),
        (
            "UNESCO",
            "Guidance for Generative AI in Education and Research",
            "2023",
            "一般原则；仅采用元数据，原始条目待再次核验",
        ),
    ]
    for record in records:
        row = table.add_row()
        for cell, value in zip(row.cells, record, strict=True):
            _set_cell_text(cell, value)

    document.add_heading("附页二：可选访谈提纲（待补材料）", level=2)
    _add_body_paragraph(
        document,
        "本报告未实施访谈。若课程教师批准并要求后续访谈，可在取得参与者知情同意、保护个人信息的前提下，"
        "围绕以下问题开展并保存真实记录：",
    )
    questions = [
        "你通常在哪些学习任务中考虑使用生成式人工智能？",
        "你如何判断工具输出是否准确，通常核对哪些来源？",
        "课程对人工智能辅助的允许范围是否清楚？",
        "使用过程中遇到过哪些隐私、诚信或能力依赖方面的顾虑？",
        "你希望学校或教师提供哪些规范与支持？",
    ]
    for index, question in enumerate(questions, start=1):
        _add_body_paragraph(document, f"{index}. {question}", indent=False)

    document.add_heading("附页三：照片粘贴位（待补材料）", level=2)
    _add_body_paragraph(
        document,
        "本报告当前没有实践照片。若后续真实开展经批准的线下活动，可在取得相关人员许可后粘贴真实照片，"
        "并填写拍摄时间、地点、活动内容和材料来源。",
    )
    photo_table = document.add_table(rows=1, cols=1)
    photo_table.style = "Table Grid"
    photo_row = photo_table.rows[0]
    photo_row.height = Cm(8)
    photo_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    _set_cell_text(photo_row.cells[0], "照片粘贴位（当前为空，不代表已完成现场实践）")


def _add_unit_evaluation(document: DocumentType) -> None:
    document.add_page_break()
    heading = document.add_heading("实践单位鉴定意见", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    warning = document.add_paragraph()
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = warning.add_run("本栏由真实实践单位填写并加盖真实公章")
    _set_run_font(run, "黑体", Pt(12), bold=True)

    table = document.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    labels = ["实践单位", "单位意见", "负责人签字", "日期", "盖章区域"]
    for row, label in zip(table.rows, labels, strict=True):
        _set_cell_text(row.cells[0], label, bold=True)
        _set_cell_text(row.cells[1], "")
    table.rows[0].height = Cm(1.2)
    table.rows[1].height = Cm(8)
    table.rows[1].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    table.rows[2].height = Cm(1.5)
    table.rows[3].height = Cm(1.5)
    table.rows[4].height = Cm(5)
    table.rows[4].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(12)
    run = note.add_run("提示：请勿自行填写单位意见、负责人签字或盖章内容。")
    _set_run_font(run, "宋体", Pt(10))


def build_report(output_path: Path) -> Path:
    output_path = Path(output_path).expanduser().resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _configure_document(document)
    _add_cover(document)
    _add_abstract(document)
    _add_main_body(document)
    _add_references(document)
    _add_process_appendix(document)
    _add_unit_evaluation(document)
    document.save(output_path)
    return output_path


def main() -> None:
    output = build_report(DEFAULT_OUTPUT_PATH)
    print(output)


if __name__ == "__main__":
    main()
