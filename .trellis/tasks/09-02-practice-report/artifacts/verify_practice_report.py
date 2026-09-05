from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn

from practice_report_content import BODY_SECTIONS, count_cjk


DEFAULT_OUTPUT_PATH = Path(r"C:\Users\Lenovo\Desktop\大学生生成式人工智能应用社会实践报告.docx")
BODY_START_MARKER = "【正文开始】"
BODY_END_MARKER = "【正文结束】"
REQUIRED_DOCUMENT_SECTIONS = [
    "摘要",
    *[heading for heading, _ in BODY_SECTIONS],
    "参考资料",
    "过程材料附页",
    "附页一：资料检索记录（现有材料）",
    "附页二：可选访谈提纲（待补材料）",
    "附页三：照片粘贴位（待补材料）",
    "实践单位鉴定意见",
]
METHOD_BOUNDARY_TERMS = [
    "公开资料分析",
    "规范文本解读",
    "典型应用场景分析",
    "未开展问卷、访谈或现场观察",
]
FORBIDDEN_FIELDWORK_CLAIMS = [
    "发放问卷",
    "有效问卷",
    "受访者表示",
    "实地走访发现",
    "调查显示",
    "访谈发现",
]


def _require(condition: bool, message: str) -> None:
    if not condition:
        raise ValueError(message)


def _body_paragraphs(paragraphs: list[str]) -> list[str]:
    _require(BODY_START_MARKER in paragraphs, "DOCX 缺少正文开始标记")
    _require(BODY_END_MARKER in paragraphs, "DOCX 缺少正文结束标记")
    start = paragraphs.index(BODY_START_MARKER)
    end = paragraphs.index(BODY_END_MARKER)
    _require(start < end, "DOCX 正文标记顺序错误")
    return paragraphs[start + 1 : end]


def _table_text(document) -> list[list[str]]:
    return [
        [cell.text.strip() for cell in row.cells]
        for table in document.tables
        for row in table.rows
    ]


def _evaluation_fields_are_blank(rows: list[list[str]]) -> bool:
    labels = {"单位意见", "负责人签字", "日期", "盖章区域"}
    values = {row[0]: row[1] for row in rows if len(row) >= 2 and row[0] in labels}
    return set(values) == labels and all(not value for value in values.values())


def _style_contract_is_valid(document, paragraph_texts: list[str]) -> bool:
    normal = document.styles["Normal"]
    heading_one = document.styles["Heading 1"]
    normal_east_asia = normal._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia"))
    heading_east_asia = heading_one._element.get_or_add_rPr().rFonts.get(qn("w:eastAsia"))
    if not (
        normal.font.size
        and abs(normal.font.size.pt - 12) <= 0.01
        and normal_east_asia == "宋体"
        and heading_one.font.size
        and abs(heading_one.font.size.pt - 16) <= 0.01
        and heading_east_asia == "黑体"
    ):
        return False

    start = paragraph_texts.index(BODY_START_MARKER)
    end = paragraph_texts.index(BODY_END_MARKER)
    body_paragraphs = [
        paragraph
        for paragraph in document.paragraphs[start + 1 : end]
        if paragraph.style.name == "Normal" and paragraph.text.strip()
    ]
    return bool(body_paragraphs) and all(
        paragraph.paragraph_format.line_spacing_rule == WD_LINE_SPACING.ONE_POINT_FIVE
        and paragraph.paragraph_format.first_line_indent
        and abs(paragraph.paragraph_format.first_line_indent.pt - 24) <= 0.01
        for paragraph in body_paragraphs
    )


def verify_report(path: Path) -> dict[str, Any]:
    resolved_path = Path(path).expanduser().resolve()
    _require(resolved_path.is_file(), f"DOCX 不存在：{resolved_path}")

    with zipfile.ZipFile(resolved_path) as archive:
        bad_member = archive.testzip()
        _require(bad_member is None, f"DOCX ZIP 成员损坏：{bad_member}")
        names = archive.namelist()
        media_names = [name for name in names if name.startswith("word/media/")]
        footer_xml = "".join(
            archive.read(name).decode("utf-8")
            for name in names
            if name.startswith("word/footer") and name.endswith(".xml")
        )

    document = Document(resolved_path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    all_text = "\n".join(paragraphs)
    body_paragraphs = _body_paragraphs(paragraphs)
    body_text = "\n".join(body_paragraphs)
    body_cjk_count = count_cjk(body_text)

    found_sections = [section for section in REQUIRED_DOCUMENT_SECTIONS if section in paragraphs]
    missing_sections = sorted(set(REQUIRED_DOCUMENT_SECTIONS) - set(found_sections))
    _require(not missing_sections, f"DOCX 缺少必需章节：{missing_sections}")
    _require(body_cjk_count >= 3000, f"正文字数不足：{body_cjk_count} 个汉字")
    _require(
        all(term in all_text for term in METHOD_BOUNDARY_TERMS),
        "DOCX 未完整说明研究方法与证据边界",
    )
    _require(
        all(term not in body_text for term in FORBIDDEN_FIELDWORK_CLAIMS),
        "正文包含可能虚构实地工作的表述",
    )
    _require(not re.search(r"\d+(?:\.\d+)?%", body_text), "正文包含未经允许的百分比表述")
    _require(len(document.inline_shapes) == 0, "DOCX 包含内嵌图片")
    _require(not media_names, f"DOCX 容器包含媒体文件：{media_names}")

    for section in document.sections:
        _require(section.orientation == WD_ORIENT.PORTRAIT, "DOCX 不是 A4 纵向")
        _require(abs(section.page_width.cm - 21) <= 0.01, "DOCX 页面宽度不是 A4")
        _require(abs(section.page_height.cm - 29.7) <= 0.01, "DOCX 页面高度不是 A4")
        _require(abs(section.top_margin.cm - 2.54) <= 0.01, "DOCX 上边距错误")
        _require(abs(section.bottom_margin.cm - 2.54) <= 0.01, "DOCX 下边距错误")
        _require(abs(section.left_margin.cm - 3) <= 0.01, "DOCX 左边距错误")
        _require(abs(section.right_margin.cm - 3) <= 0.01, "DOCX 右边距错误")

    rows = _table_text(document)
    identity_fields = ["学校", "姓名", "学号", "班级", "实践单位", "实践日期"]
    for field in identity_fields:
        _require(any(row and row[0] == field for row in rows), f"DOCX 缺少可填写字段：{field}")
    _require(_evaluation_fields_are_blank(rows), "单位鉴定意见、签字、日期或盖章区域不是空白")
    _require("PAGE" in footer_xml, "DOCX 页脚缺少页码域")
    _require(_style_contract_is_valid(document, paragraphs), "DOCX 中文字体、字号、行距或首行缩进错误")

    return {
        "path": str(resolved_path),
        "body_cjk_count": body_cjk_count,
        "required_sections": found_sections,
        "inline_shapes": len(document.inline_shapes),
        "media_files": len(media_names),
        "zip_ok": True,
        "a4_portrait": True,
        "method_boundary": True,
        "blank_evaluation_fields": True,
        "page_number_field": True,
        "style_contract": True,
    }


def main() -> None:
    print(json.dumps(verify_report(DEFAULT_OUTPUT_PATH), ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
