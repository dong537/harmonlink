from __future__ import annotations

import importlib
import re
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.enum.section import WD_ORIENT


REQUIRED_BODY_SECTIONS = [
    "一、绪论",
    "二、研究目标",
    "三、研究方法与边界",
    "四、主要应用场景",
    "五、主要风险与问题",
    "六、成因分析",
    "七、规范使用建议",
    "八、专业视角下的能力提升",
    "九、个人反思",
    "十、结语",
]


def load_content_module():
    return importlib.import_module("practice_report_content")


def test_content_model_has_required_title_sections_and_length():
    content = load_content_module()
    body = "\n".join(
        paragraph
        for _, paragraphs in content.BODY_SECTIONS
        for paragraph in paragraphs
    )

    assert content.REPORT_TITLE == "大学生生成式人工智能工具应用现状、风险与规范使用研究"
    assert [heading for heading, _ in content.BODY_SECTIONS] == REQUIRED_BODY_SECTIONS
    assert content.count_cjk(body) >= 3000


def test_content_model_states_evidence_boundary_without_fabricated_fieldwork():
    content = load_content_module()
    body = "\n".join(
        paragraph
        for _, paragraphs in content.BODY_SECTIONS
        for paragraph in paragraphs
    )

    assert "公开资料分析" in body
    assert "规范文本解读" in body
    assert "典型应用场景分析" in body
    assert "未开展问卷、访谈或现场观察" in body
    forbidden = [
        "发放问卷",
        "有效问卷",
        "受访者表示",
        "实地走访发现",
        "调查显示",
        "访谈发现",
    ]
    assert all(term not in body for term in forbidden)
    assert not re.search(r"\d+(?:\.\d+)?%", body)


def test_references_include_institution_title_year_and_link():
    content = load_content_module()

    assert len(content.REFERENCES) >= 3
    for reference in content.REFERENCES:
        assert re.search(r"20\d{2}|2018", reference)
        assert "http" in reference
    assert any("教育部" in reference for reference in content.REFERENCES)
    assert any("国家互联网信息办公室" in reference for reference in content.REFERENCES)
    assert any("UNESCO" in reference for reference in content.REFERENCES)


def test_count_cjk_counts_only_cjk_unified_ideographs():
    content = load_content_module()

    assert content.count_cjk("AI 辅助学习：生成文本 123") == 8


def load_generator_module():
    return importlib.import_module("generate_practice_report")


def test_generated_docx_contract(tmp_path: Path):
    generator = load_generator_module()
    output = generator.build_report(tmp_path / "report.docx")
    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert output.exists()
    assert "大学生生成式人工智能工具应用现状、风险与规范使用研究" in text
    assert "研究方法与边界" in text
    assert "实践单位鉴定意见" in text
    assert "本栏由真实实践单位填写并加盖真实公章" in text
    assert "正文开始" in text and "正文结束" in text
    assert len(document.inline_shapes) == 0


def test_generated_docx_uses_a4_portrait_and_expected_margins(tmp_path: Path):
    generator = load_generator_module()
    output = generator.build_report(tmp_path / "report.docx")
    document = Document(output)

    for section in document.sections:
        assert section.orientation == WD_ORIENT.PORTRAIT
        # OOXML stores dimensions in whole twips, so centimeter values round slightly.
        assert section.page_width.cm == pytest.approx(21, abs=0.01)
        assert section.page_height.cm == pytest.approx(29.7, abs=0.01)
        assert section.top_margin.cm == pytest.approx(2.54, abs=0.01)
        assert section.bottom_margin.cm == pytest.approx(2.54, abs=0.01)
        assert section.left_margin.cm == pytest.approx(3, abs=0.01)
        assert section.right_margin.cm == pytest.approx(3, abs=0.01)


def test_generated_docx_has_no_embedded_media(tmp_path: Path):
    generator = load_generator_module()
    output = generator.build_report(tmp_path / "report.docx")

    with zipfile.ZipFile(output) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
        assert media == []


def test_generated_docx_has_editable_identity_fields_and_process_appendices(tmp_path: Path):
    generator = load_generator_module()
    output = generator.build_report(tmp_path / "report.docx")
    document = Document(output)
    text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )

    for field in ["学校", "姓名", "学号", "班级", "实践单位", "实践日期"]:
        assert field in text and "________" in text
    for appendix in ["资料检索记录", "可选访谈提纲", "照片粘贴位", "现有材料", "待补材料"]:
        assert appendix in text


def test_independent_verifier_reports_document_contract(tmp_path: Path):
    generator = load_generator_module()
    verifier = importlib.import_module("verify_practice_report")
    output = generator.build_report(tmp_path / "report.docx")

    summary = verifier.verify_report(output)

    assert summary["path"] == str(output.resolve())
    assert summary["body_cjk_count"] >= 3000
    assert summary["zip_ok"] is True
    assert summary["inline_shapes"] == 0
    assert summary["media_files"] == 0
    assert summary["style_contract"] is True
    assert set(REQUIRED_BODY_SECTIONS).issubset(summary["required_sections"])
